from __future__ import annotations

import io

from docx import Document as DocxDocument

from .base import BaseParser, ParsedDocument, ParsedSection


class DOCXParser(BaseParser):
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async def parse(self, raw: bytes, *, url: str | None = None) -> ParsedDocument:
        doc = DocxDocument(io.BytesIO(raw))
        sections: list[ParsedSection] = []
        current: list[str] = []
        heading: str | None = None

        def flush() -> None:
            body = "\n".join(current).strip()
            if body:
                sections.append(
                    ParsedSection(
                        heading_path=f"/{heading}" if heading else "/",
                        heading=heading,
                        text=body,
                    )
                )
            current.clear()

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            if para.text.strip() and style.lower().startswith("heading"):
                flush()
                heading = para.text.strip()
            elif para.text.strip():
                current.append(para.text.strip())
        flush()

        title = None
        core = doc.core_properties
        if core and core.title:
            title = core.title
        return ParsedDocument(
            title=title,
            lang=None,
            content_type=self.content_type,
            sections=sections if sections else [ParsedSection("/", None, "")],
        )
